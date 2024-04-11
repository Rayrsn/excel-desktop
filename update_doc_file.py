from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from word_content import get_content
from utils.people_date import *

import os,sys

import json

def load_client_data(data):
  """
  Loads client data from a JSON.

  Args:
      json data.

  Returns:
      list: A list of dictionaries containing client data.
  """
  try:
    data = data.get("data").get("Opening_File")
    return data
  except Exception as e:
    print(f"An error occurred while loading JSON data: {e}")
    return None


def add_dict_to_doc(doc, my_list, logo_path):
    # Add a paragraph with right alignment
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Add logo
    run = paragraph.add_run()
    run.add_picture(logo_path, width=Inches(1.5))  # Adjust size as needed

    for value in my_list:
        if isinstance(value, list):
            # Create a table with one row and a number of columns equal to the number of items in the list
            if all(isinstance(i, dict) for i in value):
                table = doc.add_table(rows=1, cols=len(value) - 1)
            else:
                table = doc.add_table(rows=1, cols=len(value))
            table.style = "Table Grid"  # Optional: set the style of the table
            for i, item in enumerate(value):
                # Set each item in its own cell
                cell = table.cell(0, i)
                cell.text = str(item)
        else:
            # If the value is not a list, just add it as a new paragraph
            doc.add_paragraph(str(value))


def create_doc(doc_name, wd_var, logo_path):
    # Create a new Document
    doc = Document()
    add_dict_to_doc(doc, get_content(wd_var), logo_path)

    # Save the document
    doc.save("./Docs/" + doc_name)


import json

def gen_docs(data):
  """
  Generates Word documents for clients from a JSON file.

  Args:
        data: JSON data containing client information.

  Returns:
      bool: True if documents were generated successfully, False otherwise.
  """
  try:
    # Load client data from JSON
    client_data = load_client_data(data)
    if not client_data:
      return False

    if not os.path.exists("./Docs"):
      os.makedirs("./Docs")

    if hasattr(sys, "_MEIPASS"):
      logo_path = os.path.join(sys._MEIPASS, "bkp_logo.jpg")
    else:
      logo_path = "./bkp_logo.jpg"

    # Extract client data for each client
    # all_data = get_all_people_data(data)
    all_data = client_data
    for data in all_data:
      wd_var = {
        "email": data["Email"],
        "file_no": data["File_No"],
        "date_opened": data["Date_Opened"],
        "fee_earner": data["Fee_Earner"],  # Assuming this field exists in JSON
        "client_username": data["Clients_Surname"],
        "client_fornames": data["Clients_Forenames"],
        "client_title": data["Clients_Title"],
        "marital_status": data["Marital_Status"],
        "letters_to_home_address": data["Letters_to_Home_Address"],
        "address": data["Address"],
        "city": data["City"],
        "postcode": data["Postcode"],
        "postal_address_if_different": data["Postal_Address_if_Different"],
        "home_telephone": data["Home_Telephone"],
        "work_telephone": data["Work_Telephone"],
        "mobile_telephone": data["Mobile_Number"],
        "occupation": data["Occupation"],
        "date_of_birth": data["Date_of_Birth"],
        "ethnicity": data["Ethnicity"],
        "prison_number": data["Prison_Number"],
        "national_insurance_no": data["National_Insurance_Number"],
        "matter_type": data["Matter_Type"],
        "m_3rd_party": data["_3rd_Party"],
        "initial": data["Initial"],
      }
      name = data["Clients_Surname"]
      forname = data["Clients_Forenames"]
      doc_name = f"{name}_{forname}.docx"
      create_doc(doc_name, wd_var, logo_path)
    return True
  except Exception as e:
    print(f"An error occurred: {e}")
    return False
